"""
Layer 2 — the ATS round-trip. docs/07.

Generate → parse the document back the way an ATS would → diff against the
structured object it came from.

This is the highest-value test in the system, and the only one that tests what
actually kills applications. A CV can pass every content check, read beautifully
to a human, and still arrive at the employer as an unparseable blob with no name
on it. Nothing else catches that.

Threshold is deliberately absolute: 100% recovery of the six fields recruiters
spend ~80% of their 7.4 seconds on (docs/07 F-C). Anything less is a defect, not
a warning.

Two independent extraction paths are used so a pass is not one library's quirk:

  1. python-docx paragraph walk — how a well-behaved parser reads the document
  2. raw XML text from word/document.xml — the floor: what survives even a
     crude extractor

A field must be recoverable by BOTH. Structural hazards (tables, images,
headers/footers, text boxes, multi-column) are asserted directly against the
document XML, because those are the constructs that cause silent field loss
rather than visible breakage.
"""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from ..generate.schemas import GeneratedCV
from ..profile import Profile


@dataclass
class RoundTripResult:
    recovered: dict[str, bool] = field(default_factory=dict)
    structural: list[str] = field(default_factory=list)
    missing_detail: dict[str, str] = field(default_factory=dict)

    @property
    def passed(self) -> bool:
        return all(self.recovered.values()) and not self.structural

    def report(self) -> str:
        lines = ["PASS" if self.passed else "FAIL"]
        for k, ok in self.recovered.items():
            mark = "ok  " if ok else "LOST"
            detail = "" if ok else f"   ({self.missing_detail.get(k, '')})"
            lines.append(f"  {mark}  {k}{detail}")
        for s in self.structural:
            lines.append(f"  HAZARD  {s}")
        return "\n".join(lines)


# --------------------------------------------------------------------------- #
# Extraction — two independent paths
# --------------------------------------------------------------------------- #

def extract_via_docx(path: Path) -> str:
    """Path 1: how a well-behaved parser walks the document."""
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    # A conforming parser would also read table cells. We assert there are none,
    # but read them anyway so their absence is proven rather than assumed.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_via_xml(path: Path) -> str:
    """Path 2: the floor — raw text nodes from the document part."""
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    # <w:t> runs carry the visible text; tabs and breaks become whitespace.
    xml = re.sub(r"<w:tab[^>]*/>", " ", xml)
    xml = re.sub(r"<w:br[^>]*/>", "\n", xml)
    xml = re.sub(r"</w:p>", "\n", xml)
    texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, re.S)
    out = " ".join(texts)
    out = (out.replace("&amp;", "&").replace("&lt;", "<")
              .replace("&gt;", ">").replace("&quot;", '"').replace("&#39;", "'"))
    return out


def _structural_hazards(path: Path) -> list[str]:
    """
    The constructs that cause SILENT field loss. docs/08 §2.1.

    Checked against the package rather than the object model, because that is
    what a parser actually receives.
    """
    problems: list[str] = []
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")

        if any(n.startswith("word/header") for n in names):
            problems.append("header part present — many parsers never read it (docs/08 §1.3)")
        if any(n.startswith("word/footer") for n in names):
            problems.append("footer part present — same risk")
        if any(n.startswith("word/media/") for n in names):
            problems.append("embedded media present — images are not parseable")

    if "<w:tbl>" in xml:
        problems.append("table present — parsers flatten or drop cell contents")
    if "<w:txbxContent>" in xml or "v:textbox" in xml:
        problems.append("text box present — frequently invisible to parsers")
    if re.search(r'<w:cols[^>]*w:num="[2-9]"', xml):
        problems.append("multi-column section — the top parser-failure cause")
    if "<w:drawing>" in xml or "<w:pict>" in xml:
        problems.append("drawing/picture element present")
    return problems


# --------------------------------------------------------------------------- #
# The check
# --------------------------------------------------------------------------- #

def _norm(s: str) -> str:
    """Collapse whitespace and normalise dashes so comparison is about content."""
    s = s.replace("–", "-").replace("—", "-").replace("’", "'")
    return re.sub(r"\s+", " ", s).strip().lower()


def _present(needle: str, *haystacks: str) -> bool:
    n = _norm(needle)
    return bool(n) and all(n in _norm(h) for h in haystacks)


def roundtrip(path: Path | str, cv: GeneratedCV, profile: Profile) -> RoundTripResult:
    path = Path(path)
    r = RoundTripResult()

    via_docx = extract_via_docx(path)
    via_xml = extract_via_xml(path)
    r.structural = _structural_hazards(path)

    c = profile.contact_block

    def check(key: str, needle: str) -> None:
        ok = _present(needle, via_docx, via_xml)
        r.recovered[key] = ok
        if not ok:
            r.missing_detail[key] = f"expected {needle!r}"

    # The six fields from docs/07 F-C, plus contact details — without which the
    # application arrives anonymous however well it parses otherwise.
    check("name", c["name"])
    check("email", c["email"])
    check("phone", c["phone"])
    check("location", c["location"])

    entries = list(cv.experience) + list(cv.additional_experience)
    if not entries:
        r.recovered["experience"] = False
        r.missing_detail["experience"] = "CV has no experience entries"
    for e in entries:
        check(f"title[{e.role_id}]", e.display_title)
        check(f"employer[{e.role_id}]", e.employer)
        check(f"dates[{e.role_id}]", e.dates)

    for ed in cv.education:
        check(f"education[{ed.evidence_id}]", ed.credential)

    return r


def assert_roundtrip(path: Path | str, cv: GeneratedCV, profile: Profile) -> RoundTripResult:
    """Raise on failure. Used as a hard gate before a document may be sent."""
    res = roundtrip(path, cv, profile)
    if not res.passed:
        raise AssertionError("ATS round-trip failed:\n" + res.report())
    return res


__all__ = ["roundtrip", "assert_roundtrip", "RoundTripResult",
           "extract_via_docx", "extract_via_xml"]
