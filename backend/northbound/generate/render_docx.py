"""
Native DOCX rendering. docs/07 F-A, docs/08 §2.

DOCX is the canonical artefact — it parses at ~97% across ATS platforms against
~72% for PDF. Built natively from the structured object with python-docx; never
converted from HTML or PDF, because conversion reintroduces exactly the layout
artefacts that break parsers.

Every structural choice here is a parser constraint, not a style preference:

  no tables          parsers flatten or drop cell contents
  no text boxes      frequently invisible to parsers
  no images/icons    not parseable, often stripped
  no headers/footers many parsers never read that region — the contact block
                     would vanish and the application arrive anonymous
  single column      the top parser-failure cause
  tab stops for dates, not table cells

Layout priority follows the eye-tracking finding (docs/07 F-C): 7.4 seconds,
~80% of it on name, current title, current employer, previous title/employer,
dates and education. Those get the top of the page and the bold weight.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor

from ..profile import Profile
from .schemas import CoverLetter, ExperienceEntry, GeneratedCV

# Conservative, universally-available fonts. An exotic font is a rendering risk
# on the employer's machine and buys nothing an ATS can see.
BODY_FONT = "Calibri"
BODY_PT = 10.5
NAME_PT = 18
SECTION_PT = 11.5

RIGHT_TAB_CM = 17.0   # A4 usable width with 2cm margins


def _configure(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_PT)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.06

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(48)
        # Explicitly leave header/footer empty. docs/08 §1.3.


def _para(doc, text="", *, bold=False, size=None, space_before=0, space_after=0,
          align=None, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p


def _section_heading(doc, text: str) -> None:
    """
    Standard headings only — parsers pattern-match on these strings.
    A rule beneath, drawn as a paragraph border rather than a table or image.
    """
    p = _para(doc, text.upper(), bold=True, size=SECTION_PT, space_before=11, space_after=3)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    pPr.append(borders)


def _bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Pt(14)
    run = p.add_run(text)
    run.font.size = Pt(BODY_PT)


def _entry(doc, e: ExperienceEntry, *, compact: bool = False) -> None:
    """
    Title bold on line 1 with the date range right-aligned via a TAB STOP.

    A table would align more neatly and is exactly what breaks parsers, so the
    tab stop is deliberate.
    """
    p = _para(doc, space_before=7 if not compact else 4, space_after=0)
    p.paragraph_format.tab_stops.add_tab_stop(Pt(RIGHT_TAB_CM * 28.35 / 10), WD_TAB_ALIGNMENT.RIGHT)
    title = p.add_run(e.display_title)
    title.bold = True
    title.font.size = Pt(BODY_PT + 0.5)
    p.add_run("\t")
    dates = p.add_run(e.dates)
    dates.font.size = Pt(BODY_PT)
    dates.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Employer line — context clause included so a Canadian reader knows what
    # this company is (docs/08 §3.2).
    second = e.employer
    if e.employer_context:
        second += f" — {e.employer_context}"
    if e.location:
        second += f" · {e.location}"
    if e.employment_type:
        second += f" · {e.employment_type}"
    _para(doc, second, italic=True, size=BODY_PT - 0.5, space_after=2)

    for b in e.bullets:
        _bullet(doc, b.text)


# --------------------------------------------------------------------------- #

def render_cv(cv: GeneratedCV, profile: Profile, out_path: Path | str) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure(doc)
    c = profile.contact_block

    # ---- Contact block: BODY TEXT, first block on the page --------------- #
    # docs/08 §1.3. Never a header — many parsers never read that region.
    _para(doc, c["name"], bold=True, size=NAME_PT, space_after=2)
    _para(doc, c["location"], size=BODY_PT, space_after=1)
    line2 = "  |  ".join(x for x in (c["phone"], c["email"]) if x)
    _para(doc, line2, size=BODY_PT, space_after=1)
    links = "  |  ".join(
        x.replace("https://", "") for x in (c["linkedin"], c["portfolio"]) if x)
    if links:
        _para(doc, links, size=BODY_PT - 0.5, space_after=2)

    # ---- Summary ---------------------------------------------------------- #
    _section_heading(doc, "Professional Summary" if cv.track == "direct" else "Summary")
    _para(doc, cv.summary, space_after=1)

    # ---- Track A: skills high (keyword density for the parser) ------------ #
    if cv.track == "direct" and cv.skills:
        _section_heading(doc, "Technical Skills")
        for group, items in cv.skills.items():
            p = _para(doc, space_after=1.5)
            g = p.add_run(f"{group}: ")
            g.bold = True
            g.font.size = Pt(BODY_PT)
            r = p.add_run(", ".join(items))
            r.font.size = Pt(BODY_PT)

    # ---- Experience -------------------------------------------------------- #
    # Track B leads with the relevant physical work; the software career sits
    # under Additional Experience (docs/08 §2.4).
    _section_heading(doc, "Work Experience" if cv.track == "direct" else "Relevant Experience")
    for e in cv.experience:
        _entry(doc, e)

    # ---- Track B: skills after the relevant experience -------------------- #
    if cv.track == "transferable" and cv.skills:
        _section_heading(doc, "Skills")
        for group, items in cv.skills.items():
            p = _para(doc, space_after=1.5)
            g = p.add_run(f"{group}: ")
            g.bold = True
            g.font.size = Pt(BODY_PT)
            r = p.add_run(", ".join(items))
            r.font.size = Pt(BODY_PT)

    if cv.additional_experience:
        _section_heading(doc, "Additional Experience")
        for e in cv.additional_experience:
            _entry(doc, e, compact=True)

    # ---- Education -------------------------------------------------------- #
    _section_heading(doc, "Education" if cv.track == "direct" else "Education & Training")
    for ed in cv.education:
        p = _para(doc, space_before=4, space_after=0)
        p.paragraph_format.tab_stops.add_tab_stop(
            Pt(RIGHT_TAB_CM * 28.35 / 10), WD_TAB_ALIGNMENT.RIGHT)
        run = p.add_run(ed.credential)
        run.bold = True
        run.font.size = Pt(BODY_PT + 0.5)
        p.add_run("\t")
        yr = p.add_run(ed.year)
        yr.font.size = Pt(BODY_PT)
        yr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        _para(doc, ed.institution, italic=True, size=BODY_PT - 0.5, space_after=1)
        if ed.detail:
            _para(doc, ed.detail, size=BODY_PT, space_after=1)

    # ---- Portfolio (Track A only) ----------------------------------------- #
    if cv.track == "direct" and cv.portfolio_ids:
        _section_heading(doc, "Projects")
        for pid in cv.portfolio_ids:
            ev = profile.evidence.get(pid)
            if ev:
                _bullet(doc, ev.text)

    # ---- Languages -------------------------------------------------------- #
    if cv.languages:
        _section_heading(doc, "Languages")
        _para(doc, ", ".join(cv.languages), space_after=1)

    if cv.availability:
        _section_heading(doc, "Availability")
        _para(doc, cv.availability, space_after=1)

    doc.save(str(out_path))
    return out_path


def render_cover_letter(letter: CoverLetter, profile: Profile,
                        employer: str, out_path: Path | str) -> Path:
    """
    One page, business-letter form. docs/08 §5.

    Paragraph order is fixed and paragraph 4 is the ONLY place work-permit
    status appears — the audit enforces both.
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure(doc)
    c = profile.contact_block

    _para(doc, c["name"], bold=True, size=13, space_after=1)
    _para(doc, c["location"], space_after=1)
    _para(doc, "  |  ".join(x for x in (c["phone"], c["email"]) if x), space_after=10)

    _para(doc, employer, space_after=10)
    _para(doc, letter.greeting, space_after=8)

    for para in (letter.opening, letter.evidence, letter.bridge):
        _para(doc, para, space_after=8)

    # Screening answers before the authorisation paragraph — ~30% of LMIA
    # postings ask, and most applicants never answer them (docs/07).
    if letter.screening_answers:
        for ans in letter.screening_answers:
            _para(doc, ans, space_after=4)
        _para(doc, "", space_after=4)

    _para(doc, letter.authorisation, space_after=10)
    _para(doc, letter.signoff, space_after=2)
    _para(doc, c["name"], space_after=0)

    doc.save(str(out_path))
    return out_path


__all__ = ["render_cv", "render_cover_letter"]
