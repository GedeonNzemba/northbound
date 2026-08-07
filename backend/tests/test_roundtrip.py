"""
Layer 2 tests — the ATS round-trip.

These prove the renderer produces documents an ATS can actually read. The
negative tests matter as much as the positive one: they prove the check would
catch a regression rather than passing everything.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from fixtures import PROFILE, full_cv as _track_b_cv

from northbound.evaluate.ats_roundtrip import (
    extract_via_docx, extract_via_xml, roundtrip,
)
from northbound.generate.render_docx import render_cv


@pytest.fixture(scope="module")
def rendered(tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("cv") / "cv.docx"
    return render_cv(_track_b_cv(), PROFILE, out)


def test_roundtrip_passes(rendered):
    res = roundtrip(rendered, _track_b_cv(), PROFILE)
    assert res.passed, res.report()


def test_all_six_recruiter_fields_recovered(rendered):
    """docs/07 F-C — the fields ~80% of the 7.4 seconds goes to."""
    res = roundtrip(rendered, _track_b_cv(), PROFILE)
    for key, ok in res.recovered.items():
        assert ok, f"{key} was not recoverable: {res.missing_detail.get(key)}"


def test_both_extraction_paths_agree(rendered):
    """A pass must not depend on one library's behaviour."""
    a, b = extract_via_docx(rendered), extract_via_xml(rendered)
    for needle in ("Gedeon Christ Nzemba", "gedeon@gedeonchrist.com",
                   "Cumpsty Electrical", "Val de Vie"):
        assert needle in a, f"{needle!r} missing from python-docx extraction"
        assert needle in b, f"{needle!r} missing from raw XML extraction"


def test_no_structural_hazards(rendered):
    """No tables, images, headers, footers, text boxes or columns."""
    res = roundtrip(rendered, _track_b_cv(), PROFILE)
    assert res.structural == [], res.structural


def test_contact_block_is_body_text_not_header(rendered):
    """
    docs/08 §1.3 — the single most damaging layout mistake. Contact details in
    a header are invisible to many parsers and the application arrives anonymous.
    """
    doc = Document(str(rendered))
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "Gedeon Christ Nzemba" in body
    assert "gedeon@gedeonchrist.com" in body
    import zipfile
    with zipfile.ZipFile(rendered) as z:
        assert not [n for n in z.namelist() if n.startswith("word/header")]


def test_no_street_address_or_postal_code(rendered):
    """docs/08 §1.3 — city + country only."""
    text = extract_via_xml(rendered)
    import re
    assert not re.search(r"\b\d+\s+[A-Z][a-z]+\s+(Road|Street|Drive|Flat)\b", text)
    assert not re.search(r"\b[A-Z]\d[A-Z]\s?\d[A-Z]\d\b", text)


def test_no_day_level_dates(rendered):
    """The painter/packer resolution must survive rendering."""
    import re
    text = extract_via_xml(rendered)
    assert not re.search(
        r"\b\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b",
        text), "day-level date survived into the rendered document"


def test_track_b_puts_software_below_relevant_experience(rendered):
    """docs/08 §2.4 — a farm employer must hit physical work first."""
    text = extract_via_docx(rendered)
    assert text.index("RELEVANT EXPERIENCE") < text.index("ADDITIONAL EXPERIENCE")
    assert text.index("Cumpsty Electrical") < text.index("Kurtosys Systems")


def test_icas_equivalency_survives_rendering(rendered):
    """docs/08 §3.1 — most overseas applicants never state this."""
    text = extract_via_xml(rendered)
    assert "ICAS" in text and "equivalent to Canadian Secondary School Graduation" in text
    assert "24080341" in text


# ---- layout geometry: what no text check can see -------------------------- #

def test_date_tab_stops_sit_at_the_usable_text_width(rendered):
    """
    The dates column right-aligns at the margin, and nothing in the text layer
    can tell you whether it does.

    This is a regression test for a shipped bug: the tab stop was computed from
    a hardcoded 17cm constant divided by ten, putting it at 1.7cm. Every title
    is longer than that, so the tab fell through to a default stop and the date
    landed just after the title with an arbitrary gap. The text extracted
    identically, every round-trip check passed, and the document a human opened
    had no aligned date column — in one of the six fields the 7.4-second scan
    actually lands on (docs/07 F-C).

    Measuring the section makes it self-correcting; asserting the measurement
    keeps it that way.
    """
    import re
    import zipfile

    doc = Document(str(rendered))
    s = doc.sections[0]
    usable_emu = int(s.page_width) - int(s.left_margin) - int(s.right_margin)
    expected_twips = round(usable_emu / 635)

    xml = zipfile.ZipFile(rendered).read("word/document.xml").decode()
    stops = [int(t) for t in re.findall(r'<w:tab w:pos="(\d+)" w:val="right"', xml)]

    assert stops, "no right tab stops — the date column is not aligned at all"
    for pos in stops:
        assert abs(pos - expected_twips) <= 2, (
            f"right tab at {pos} twips ({pos / 20 / 28.35:.2f} cm), expected "
            f"{expected_twips} ({expected_twips / 20 / 28.35:.2f} cm — the margin)")


def test_every_dated_line_has_a_tab_stop(rendered):
    """One per experience entry and one per education entry, or a date is adrift."""
    import re
    import zipfile

    xml = zipfile.ZipFile(rendered).read("word/document.xml").decode()
    stops = len(re.findall(r'<w:tab w:pos="\d+" w:val="right"', xml))
    cv = _track_b_cv()
    expected = len(cv.experience) + len(cv.additional_experience) + len(cv.education)
    assert stops == expected, f"{stops} tab stops for {expected} dated lines"


# ---- negative tests: prove the check would catch a regression ------------- #

def test_roundtrip_detects_a_table(tmp_path):
    doc = Document()
    doc.add_paragraph("Gedeon Christ Nzemba")
    doc.add_table(rows=1, cols=2)
    p = tmp_path / "bad.docx"
    doc.save(str(p))
    res = roundtrip(p, _track_b_cv(), PROFILE)
    assert not res.passed
    assert any("table" in s for s in res.structural)


def test_roundtrip_detects_contact_details_in_a_header(tmp_path):
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Gedeon Christ Nzemba | gedeon@gedeonchrist.com"
    doc.add_paragraph("Work Experience")
    p = tmp_path / "header.docx"
    doc.save(str(p))
    res = roundtrip(p, _track_b_cv(), PROFILE)
    assert not res.passed
    assert any("header" in s for s in res.structural)
    assert res.recovered["name"] is False, "name in a header must count as lost"
